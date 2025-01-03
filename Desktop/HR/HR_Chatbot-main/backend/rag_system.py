# rag_system.py

import os
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
from datetime import datetime
import google.generativeai as genai
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import logging
from pathlib import Path

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class DocumentMetadata:
    """Metadata for processed documents"""
    title: str
    file_type: str
    created_date: str
    collection_id: str
    page_count: int
    status: str  # 'processing', 'active', 'error'
    error_message: Optional[str] = None
    last_updated: str = datetime.now().isoformat()

@dataclass
class ProcessedDocument:
    """Processed document information"""
    chunks: List[str]
    metadata: DocumentMetadata
    collection_path: str
    total_chunks: int
    embedding_status: str  # 'pending', 'completed', 'error'

class DocumentProcessor:
    """Handles document processing and embedding"""
    
    def __init__(self, api_key: str, base_path: str = "./data"):
        """
        Initialize the document processor.
        
        Args:
            api_key: Google API key for Gemini
            base_path: Base path for storing documents and collections
        """
        self.api_key = api_key
        self.base_path = base_path
        genai.configure(api_key=api_key)
        
        # Initialize embeddings
        self.embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=api_key
        )
        
        # Configure text splitter for Arabic and English
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            length_function=len,
            separators=["\n\n", "\n", ".", "!", "?", "؟", "،", " "]
        )
        
        # Ensure directory structure exists
        self.docs_path = Path(base_path) / "documents"
        self.chroma_path = Path(base_path) / "chroma_db"
        self.docs_path.mkdir(parents=True, exist_ok=True)
        self.chroma_path.mkdir(parents=True, exist_ok=True)

    def process_document(self, file_path: str) -> ProcessedDocument:
        """
        Process a document and prepare it for embedding.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ProcessedDocument object containing chunks and metadata
        """
        try:
            logger.info(f"Starting document processing: {file_path}")
            
            # Extract text based on file type
            text = self._extract_text(file_path)
            
            # Create document metadata
            metadata = self._create_metadata(file_path)
            
            # Split text into chunks
            chunks = self.text_splitter.split_text(text)
            
            # Create collection path
            collection_path = self.chroma_path / metadata.collection_id
            
            # Create processed document
            processed_doc = ProcessedDocument(
                chunks=chunks,
                metadata=metadata,
                collection_path=str(collection_path),
                total_chunks=len(chunks),
                embedding_status='pending'
            )
            
            logger.info(f"Document processed successfully: {metadata.collection_id}")
            return processed_doc
            
        except Exception as e:
            error_msg = f"Error processing document: {str(e)}"
            logger.error(error_msg)
            raise Exception(error_msg)

    def _extract_text(self, file_path: str) -> str:
        """Extract text from different document types."""
        file_type = Path(file_path).suffix.lower()
        
        try:
            if file_type == '.pdf':
                return self._extract_pdf_text(file_path)
            elif file_type in ['.docx', '.doc']:
                return self._extract_word_text(file_path)
            elif file_type == '.txt':
                return self._extract_txt_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            raise Exception(f"Error extracting text: {str(e)}")

    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF files."""
        try:
            import pypdf
            
            with open(file_path, 'rb') as file:
                pdf = pypdf.PdfReader(file)
                text = ""
                for page in pdf.pages:
                    text += page.extract_text() + "\n\n"
                return text
                
        except Exception as e:
            raise Exception(f"Error extracting PDF text: {str(e)}")

    def _extract_word_text(self, file_path: str) -> str:
        """Extract text from Word documents."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
            
        except Exception as e:
            raise Exception(f"Error extracting Word text: {str(e)}")

    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from txt files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"Error extracting text file text: {str(e)}")

    def _create_metadata(self, file_path: str) -> DocumentMetadata:
        """Create metadata for a document."""
        file_path = Path(file_path)
        creation_time = datetime.fromtimestamp(file_path.stat().st_ctime)
        
        return DocumentMetadata(
            title=file_path.name,
            file_type=file_path.suffix[1:],  # Remove the dot
            created_date=creation_time.isoformat(),
            collection_id=f"doc_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            page_count=0,  # Will be updated based on document type
            status='processing'
        )

    def create_embeddings(self, processed_doc: ProcessedDocument) -> Chroma:
        """Create embeddings for processed document chunks."""
        try:
            logger.info(f"Creating embeddings for document: {processed_doc.metadata.collection_id}")
            
            # Prepare metadata for each chunk
            chunk_metadata = []
            for i, chunk in enumerate(processed_doc.chunks):
                chunk_metadata.append({
                    "chunk_id": str(i),
                    "document_title": processed_doc.metadata.title,
                    "collection_id": processed_doc.metadata.collection_id
                })
            
            # Create and persist vector store
            vectorstore = Chroma.from_texts(
                texts=processed_doc.chunks,
                metadatas=chunk_metadata,
                embedding=self.embeddings,
                persist_directory=processed_doc.collection_path
            )
            
            logger.info(f"Embeddings created successfully: {processed_doc.metadata.collection_id}")
            return vectorstore
            
        except Exception as e:
            error_msg = f"Error creating embeddings: {str(e)}"
            logger.error(error_msg)
            raise Exception(error_msg)
        
        
# Continuing rag_system.py

class HRRAGSystem:
    """Main RAG system for HR document queries"""
    
    def __init__(self, api_key: str, base_path: str = "./data"):
        """
        Initialize the HR RAG system.
        
        Args:
            api_key: Google API key for Gemini
            base_path: Base path for data storage
        """
        self.api_key = api_key
        self.base_path = Path(base_path)
        self.doc_processor = DocumentProcessor(api_key, base_path)
        
        # Initialize LLM
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-pro",
            google_api_key=api_key,
            temperature=0.3,
            convert_system_message_to_human=True
        )
        
        # Keep track of active collections
        self.active_collections: Dict[str, Chroma] = {}
        
        # Initialize from existing collections if any
        self._load_existing_collections()

    def _load_existing_collections(self):
        """Load existing collections from chroma_db directory."""
        try:
            chroma_path = self.base_path / "chroma_db"
            if chroma_path.exists():
                for collection_dir in chroma_path.iterdir():
                    if collection_dir.is_dir():
                        try:
                            vectorstore = Chroma(
                                persist_directory=str(collection_dir),
                                embedding_function=self.doc_processor.embeddings
                            )
                            self.active_collections[collection_dir.name] = vectorstore
                            logger.info(f"Loaded existing collection: {collection_dir.name}")
                        except Exception as e:
                            logger.error(f"Error loading collection {collection_dir.name}: {str(e)}")
        
        except Exception as e:
            logger.error(f"Error loading existing collections: {str(e)}")

    def process_document(self, file_path: str) -> str:
        """
        Process a new document and add it to the RAG system.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            collection_id: ID of the created collection
        """
        try:
            # Process document
            processed_doc = self.doc_processor.process_document(file_path)
            
            # Create embeddings
            vectorstore = self.doc_processor.create_embeddings(processed_doc)
            
            # Add to active collections
            self.active_collections[processed_doc.metadata.collection_id] = vectorstore
            
            # Update metadata status
            processed_doc.metadata.status = 'active'
            processed_doc.embedding_status = 'completed'
            
            return processed_doc.metadata.collection_id
            
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            raise

    def query(self, question: str, collection_ids: Optional[List[str]] = None) -> Dict[str, Any]:
        """
        Query the RAG system with a question.
        
        Args:
            question: The question to ask
            collection_ids: Optional list of specific collection IDs to query
            
        Returns:
            Dict containing answer and source documents
        """
        try:
            # If no specific collections provided, use all active ones
            collections_to_query = []
            if collection_ids:
                for cid in collection_ids:
                    if cid in self.active_collections:
                        collections_to_query.append(self.active_collections[cid])
            else:
                collections_to_query = list(self.active_collections.values())

            if not collections_to_query:
                return {
                    "answer": "عذراً، لا توجد مستندات متاحة للبحث.",
                    "source_documents": []
                }

            # Combine results from all collections
            all_docs = []
            for vectorstore in collections_to_query:
                docs = vectorstore.similarity_search(question, k=2)
                all_docs.extend(docs)

            # Sort by relevance (assumed from order) and take top results
            context = "\n".join(doc.page_content for doc in all_docs[:3])

            # Generate response using LLM
            prompt = f"""أنت مساعد متخصص في الموارد البشرية. استخدم المعلومات التالية للإجابة على السؤال.
            إذا لم تجد المعلومات في النص المتوفر، قل ذلك بصراحة.

            السؤال: {question}

            السياق:
            {context}

            الإجابة:"""

            response = self.llm.invoke(prompt)

            return {
                "answer": response.content,
                "source_documents": all_docs[:3]
            }

        except Exception as e:
            logger.error(f"Error processing query: {str(e)}")
            return {
                "answer": "عذراً، حدث خطأ أثناء معالجة السؤال.",
                "source_documents": []
            }

    def merge_collections(self, collection_ids: List[str]) -> Optional[str]:
        """
        Merge multiple collections into a new one.
        
        Args:
            collection_ids: List of collection IDs to merge
            
        Returns:
            New collection ID if successful, None otherwise
        """
        try:
            # Validate collections exist
            collections = []
            for cid in collection_ids:
                if cid in self.active_collections:
                    collections.append(self.active_collections[cid])
                else:
                    raise ValueError(f"Collection not found: {cid}")

            # Create new collection ID
            merged_id = f"merged_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            merged_path = self.base_path / "chroma_db" / merged_id

            # Get all documents from collections
            all_texts = []
            all_metadata = []
            
            for vectorstore in collections:
                docs = vectorstore.get()
                all_texts.extend(docs['documents'])
                all_metadata.extend(docs['metadatas'])

            # Create new merged collection
            merged_vectorstore = Chroma.from_texts(
                texts=all_texts,
                metadatas=all_metadata,
                embedding=self.doc_processor.embeddings,
                persist_directory=str(merged_path)
            )

            # Add to active collections
            self.active_collections[merged_id] = merged_vectorstore

            return merged_id

        except Exception as e:
            logger.error(f"Error merging collections: {str(e)}")
            return None

    def remove_collection(self, collection_id: str) -> bool:
        """
        Remove a collection from the system.
        
        Args:
            collection_id: ID of the collection to remove
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            if collection_id in self.active_collections:
                # Remove from active collections
                del self.active_collections[collection_id]
                
                # Remove directory
                collection_path = self.base_path / "chroma_db" / collection_id
                if collection_path.exists():
                    import shutil
                    shutil.rmtree(str(collection_path))
                
                logger.info(f"Removed collection: {collection_id}")
                return True
                
            return False
            
        except Exception as e:
            logger.error(f"Error removing collection: {str(e)}")
            return False

    def get_active_collections(self) -> List[Dict[str, Any]]:
        """
        Get information about all active collections.
        
        Returns:
            List of collection information dictionaries
        """
        try:
            collections_info = []
            for cid, vectorstore in self.active_collections.items():
                try:
                    docs = vectorstore.get()
                    collections_info.append({
                        "collection_id": cid,
                        "document_count": len(docs['documents']),
                        "created": datetime.fromtimestamp(
                            (self.base_path / "chroma_db" / cid).stat().st_ctime
                        ).isoformat()
                    })
                except Exception as e:
                    logger.error(f"Error getting info for collection {cid}: {str(e)}")
                    
            return collections_info
            
        except Exception as e:
            logger.error(f"Error getting collections info: {str(e)}")
            return []